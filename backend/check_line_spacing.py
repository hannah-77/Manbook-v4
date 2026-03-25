import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document()
s = doc.styles['Normal']
s.paragraph_format.line_spacing = Pt(16.5)

crop_local_val = 'c:\\Users\\Hanna\\Manbook-v4\\backend\\output_results\\101. rev01 - PTB 2in1 eng.docx_0_crop_figure_11.png'

tbl = doc.add_table(rows=1, cols=1)
tbl.style = 'Table Grid'
cell = tbl.cell(0, 0)
cell_para = cell.paragraphs[0]
cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
cell_para.paragraph_format.first_line_indent = Pt(0)
# DO NOT reset line spacing to see if it clips

run_img = cell_para.add_run()
run_img.add_picture(crop_local_val, width=Inches(4.5))

doc.save('test_bug.docx')

doc2 = Document()
s2 = doc2.styles['Normal']
s2.paragraph_format.line_spacing = Pt(16.5)

tbl2 = doc2.add_table(rows=1, cols=1)
tbl2.style = 'Table Grid'
cell2 = tbl2.cell(0, 0)
cell_para2 = cell2.paragraphs[0]
cell_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
cell_para2.paragraph_format.first_line_indent = Pt(0)

# Reset line spacing for the image paragraph
cell_para2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
cell_para2.paragraph_format.line_spacing = Pt(0) 

run_img2 = cell_para2.add_run()
run_img2.add_picture(crop_local_val, width=Inches(4.5))

doc2.save('test_fix.docx')
