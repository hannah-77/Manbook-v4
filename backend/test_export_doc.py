import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.add_heading('Test Document', level=1)

crop_local_val = 'c:\\Users\\Hanna\\Manbook-v4\\backend\\output_results\\101. rev01 - PTB 2in1 eng.docx_0_crop_figure_11.png'
print(f"File exists: {os.path.exists(crop_local_val)}")

tbl = doc.add_table(rows=1, cols=1)
tbl.style = 'Table Grid'
cell = tbl.cell(0, 0)
cell_para = cell.paragraphs[0]
cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
cell_para.paragraph_format.first_line_indent = Pt(0)
run_img = cell_para.add_run()

try:
    run_img.add_picture(crop_local_val, width=Inches(4.5))
    print("Added picture successfully")
except Exception as e:
    print(f"Exception: {e}")

doc.save('test_export_doc.docx')
print("Test doc created: test_export_doc.docx")
