"""
MODULE 3: THE ARCHITECT 🏗️ (Builder)
======================================
Bertugas menyusun kembali data ke Template Standar (.docx).
Template profesional Manual Book standar industri.

Updated: February 2026
"""

// jjj

import os
import re
import logging
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

from bio_brain import BioBrain

load_dotenv()

logger = logging.getLogger("BioManual")


def _get_base_path():
    """Get base path for resolving file locations."""
    import sys
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


def _get_output_dir():
    """Get the output_results directory path."""
    return os.path.join(_get_base_path(), "output_results")


class BioArchitect:
    """
    Bertugas menyusun kembali data ke Template Standar (.docx).
    Template profesional Manual Book standar industri.
    """

    # Warna tema buku manual (biru tua Biosys)
    COLOR_PRIMARY   = RGBColor(0x1E, 0x3A, 0x8A)   # Biru tua #1E3A8A
    COLOR_SECONDARY = RGBColor(0x3B, 0x82, 0xF6)   # Biru muda #3B82F6
    COLOR_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
    COLOR_GRAY      = RGBColor(0x6B, 0x72, 0x80)
    COLOR_BLACK     = RGBColor(0x00, 0x00, 0x00)

    def __init__(self):
        self.base_path = _get_base_path()


    def _resolve_crop_path(self, item):
        """
        Resolve the local file path for a figure/table crop image.

        Tries multiple strategies (in order):
          1. crop_local: absolute path  → if it exists on disk
          2. crop_url:   extract filename from URL → check output_results/
          3. crop_local: relative path  → resolve against base_path
          4. crop_local: basename only  → check output_results/

        Returns the resolved absolute path, or None if not found.
        """
        output_dir = _get_output_dir()

        # ── Strategy 1: crop_local as-is (absolute path) ──
        crop_local = item.get('crop_local')
        if crop_local and os.path.exists(crop_local):
            return crop_local

        # ── Strategy 2: Extract filename from crop_url ──
        crop_url = item.get('crop_url')
        if crop_url and isinstance(crop_url, str):
            try:
                from urllib.parse import unquote, urlparse
                parsed = urlparse(crop_url)
                url_filename = unquote(parsed.path.split('/')[-1])
                if url_filename:
                    candidate = os.path.join(output_dir, url_filename)
                    if os.path.exists(candidate):
                        logger.info(f"   🔄 Fallback: resolved via crop_url → {candidate}")
                        return candidate
            except Exception:
                pass

        # ── Strategy 3: crop_local as relative path ──
        if crop_local:
            candidate = os.path.join(self.base_path, crop_local)
            if os.path.exists(candidate):
                logger.info(f"   🔄 Fallback: resolved as relative path → {candidate}")
                return candidate

            # ── Strategy 4: Just the basename in output_results ──
            basename = os.path.basename(crop_local)
            if basename:
                candidate = os.path.join(output_dir, basename)
                if os.path.exists(candidate):
                    logger.info(f"   🔄 Fallback: resolved via basename → {candidate}")
                    return candidate

        return None

    # ─────────────────────────────────────────────
    # Layout & Style
    # ─────────────────────────────────────────────
    def _set_fixed_margins(self, doc):
        """Ukuran A4, margin standar manual book Indonesia."""
        for section in doc.sections:
            section.page_height  = Cm(29.7)
            section.page_width   = Cm(21.0)
            section.top_margin   = Cm(2.5)
            section.bottom_margin= Cm(2.5)
            section.left_margin  = Cm(3.0)   # lebih lebar untuk binding
            section.right_margin = Cm(2.0)
            section.header_distance = Cm(1.25)
            section.footer_distance = Cm(1.25)

    def _set_fixed_styles(self, doc):
        """Font, spacing, dan heading styles."""
        # ── Normal (Body Text) ──────────────────────────────
        s = doc.styles['Normal']
        s.font.name = 'Arial'
        s.font.size = Pt(11)
        s.font.color.rgb = self.COLOR_BLACK
        s.paragraph_format.line_spacing     = Pt(16.5)   # ≈ 1.5 × 11pt
        s.paragraph_format.space_before     = Pt(0)
        s.paragraph_format.space_after      = Pt(8)
        s.paragraph_format.first_line_indent= Inches(0.35)

        # ── Heading 1 (Judul BAB) ───────────────────────────
        h1 = doc.styles['Heading 1']
        h1.font.name = 'Arial'
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.font.color.rgb = self.COLOR_WHITE
        h1.paragraph_format.space_before    = Pt(0)
        h1.paragraph_format.space_after     = Pt(12)
        h1.paragraph_format.keep_with_next  = True
        h1.paragraph_format.first_line_indent = Pt(0)

        # ── Heading 2 (Sub-judul / 1.1) ─────────────────────────────
        h2 = doc.styles['Heading 2']
        h2.font.name = 'Arial'
        h2.font.size = Pt(12)
        h2.font.bold = True
        h2.font.color.rgb = self.COLOR_PRIMARY
        h2.paragraph_format.space_before    = Pt(10)
        h2.paragraph_format.space_after     = Pt(4)
        h2.paragraph_format.first_line_indent = Pt(0)

        # ── Heading 3 (Sub-sub-judul / 1.1.1) ───────────────────────
        h3 = doc.styles['Heading 3']
        h3.font.name = 'Arial'
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.italic = True
        h3.font.color.rgb = self.COLOR_SECONDARY
        h3.paragraph_format.space_before    = Pt(6)
        h3.paragraph_format.space_after     = Pt(2)
        h3.paragraph_format.left_indent     = Inches(0.25)  # indentasi H3
        h3.paragraph_format.first_line_indent = Pt(0)

        # ── Heading 4 (Level 4+, misal 1.1.1.1) ─────────────────────
        h4 = doc.styles['Heading 4']
        h4.font.name = 'Arial'
        h4.font.size = Pt(12)
        h4.font.bold = False
        h4.font.italic = True
        h4.font.color.rgb = self.COLOR_GRAY
        h4.paragraph_format.space_before    = Pt(4)
        h4.paragraph_format.space_after     = Pt(2)
        h4.paragraph_format.left_indent     = Inches(0.5)   # indentasi H4
        h4.paragraph_format.first_line_indent = Pt(0)

    # ─────────────────────────────────────────────
    # Header & Footer
    # ─────────────────────────────────────────────
    def _add_header_footer(self, doc, bab_label=""):
        """
        Header: KOSONG (tidak ada teks)
        Footer: letterhead.png full-width (gelombang biru + logo Elitech) + nomor halaman
        """
        section = doc.sections[0]

        # ── HEADER: Kosongkan ──────────────────────────────────────
        section.different_first_page_header_footer = True
        header = section.header
        header.is_linked_to_previous = False

        if not header.paragraphs:
            header.add_paragraph()
        hp = header.paragraphs[0]
        hp.clear()
        # Header dibiarkan kosong sesuai permintaan user

        # ── FOOTER ──────────────────────────────────────────
        footer = section.footer
        footer.is_linked_to_previous = False

        if not footer.paragraphs:
            footer.add_paragraph()
        fp = footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)

        # Tambahkan nomor halaman di atas letterhead
        run_pg = fp.add_run()
        run_pg.font.name  = 'Arial'
        run_pg.font.size  = Pt(9)
        run_pg.font.color.rgb = self.COLOR_PRIMARY
        run_pg.font.bold  = True

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = ' PAGE '
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run_pg._r.extend([fldChar1, instrText, fldChar2])

        # Standard text footer
        run_co = fp.add_run("\n© Elitech Technovision")
        run_co.font.name  = 'Arial'
        run_co.font.size  = Pt(8)
        run_co.font.color.rgb = self.COLOR_GRAY

    # ─────────────────────────────────────────────
    # Cover Page
    # ─────────────────────────────────────────────
    def _build_cover_page(self, doc, original_filename, product_name="", product_desc="", lang='id'):
        """
        Cover page layout:
          - Kiri atas  : Nama produk bold + deskripsi
          - Tengah     : "BUKU MANUAL" atau "MANUAL BOOK" bold besar
          - Bawah      : Letterhead full-width (wave + logo)
        """
        # Extract product_name from original_filename if missing or too messy
        if not product_name:
            stem = Path(original_filename).stem
            m = re.search(r'(?i)(?:buku manual|manual book|user manual)\s+([A-Z0-9\-]+)', stem)
            if m and m.group(1):
                product_name = m.group(1)
            else:
                product_name = stem.split('.')[0]
                if len(product_name) > 20:
                    product_name = product_name[:20].strip()

        # ── Kiri atas: Nama Produk ───────────────────────────
        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_name.paragraph_format.space_before = Pt(60)
        p_name.paragraph_format.space_after  = Pt(0)
        p_name.paragraph_format.first_line_indent = Pt(0)
        rn = p_name.add_run(product_name.upper())
        rn.font.name  = 'Arial'
        rn.font.size  = Pt(44)
        rn.font.bold  = True
        rn.font.color.rgb = self.COLOR_BLACK

        # Deskripsi produk (di bawah nama produk)
        if product_desc:
            p_desc = doc.add_paragraph()
            p_desc.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_desc.paragraph_format.space_before = Pt(8)
            p_desc.paragraph_format.space_after  = Pt(0)
            p_desc.paragraph_format.first_line_indent = Pt(0)
            rd = p_desc.add_run(product_desc)
            rd.font.name  = 'Arial'
            rd.font.size  = Pt(14)
            rd.font.bold  = False
            rd.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # ── Spasi tengah (mendorong MANUAL BOOK ke bawah) ────
        for _ in range(8):
            sp = doc.add_paragraph()
            sp.paragraph_format.first_line_indent = Pt(0)
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after  = Pt(0)

        # ── Kanan bawah: BUKU MANUAL / MANUAL BOOK ──────────
        p_bm = doc.add_paragraph()
        p_bm.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_bm.paragraph_format.space_before = Pt(80)
        p_bm.paragraph_format.space_after = Pt(0)
        p_bm.paragraph_format.first_line_indent = Pt(0)
        p_bm.paragraph_format.right_indent = Pt(40)
        
        rbm = p_bm.add_run("MANUAL BOOK" if lang == 'en' else "BUKU MANUAL")
        rbm.font.name  = 'Arial'
        rbm.font.size  = Pt(44)
        rbm.font.bold  = True
        rbm.font.color.rgb = self.COLOR_BLACK

        # No letterhead branding on cover
        p_co = doc.add_paragraph()
        p_co.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_co = p_co.add_run("© Elitech Technovision")
        run_co.font.name  = 'Arial'
        run_co.font.size  = Pt(10)
        run_co.font.color.rgb = self.COLOR_GRAY
        
        doc.add_page_break()

    # ─────────────────────────────────────────────
    # Chapter Header (kotak biru)
    # ─────────────────────────────────────────────
    def _add_chapter_header(self, doc, bab_id, bab_title):
        """Judul BAB dengan kotak biru tua + teks putih."""
        h = doc.add_heading(f"  {bab_id}: {bab_title}", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h.paragraph_format.first_line_indent = Pt(0)

        # Shading biru pada paragraf heading
        pPr = h._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '1E3A8A')   # biru tua
        pPr.append(shd)

        # Padding atas-bawah via spacing
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after  = Pt(14)

    # ─────────────────────────────────────────────
    # Daftar Isi (Table of Contents)
    # ─────────────────────────────────────────────
    def _build_toc_page(self, doc, grouped, lang='id'):
        """
        Halaman Daftar Isi.
        - Insert field TOC standar Word (auto-update saat buka file)
        - Fallback: daftar statis BAB 1-7 dengan nomor halaman estimasi
        """
        # ── Judul "DAFTAR ISI" ─────────────────────────────
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after  = Pt(20)
        p_title.paragraph_format.first_line_indent = Pt(0)
        rt = p_title.add_run("TABLE OF CONTENTS" if lang == 'en' else "DAFTAR ISI")
        rt.font.name  = 'Arial'
        rt.font.size  = Pt(16)
        rt.font.bold  = True
        rt.font.color.rgb = self.COLOR_PRIMARY

        # Garis bawah biru
        pPr = p_title._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),  'single')
        bot.set(qn('w:sz'),   '8')
        bot.set(qn('w:space'),'2')
        bot.set(qn('w:color'), '1E3A8A')
        pBdr.append(bot)
        pPr.append(pBdr)

        doc.add_paragraph().paragraph_format.first_line_indent = Pt(0)

        # ── Insert TOC field (Word akan render otomatis) ────────
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(0)
        run = paragraph.add_run()

        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-2" \\h \\z \\u '

        fldChar_sep = OxmlElement('w:fldChar')
        fldChar_sep.set(qn('w:fldCharType'), 'separate')

        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')

        run._r.extend([fldChar_begin, instrText, fldChar_sep, fldChar_end])

        # ── Fallback statis: daftar BAB yang ada ──────────────
        bab_meta = BioBrain().taxonomy
        page_est  = 3   # estimasi mulai halaman 3 (setelah cover + toc)

        for bab_id, items in grouped.items():
            # Skip English chapters if Indonesian, skip Indonesian if English
            if lang == 'en' and bab_id.startswith('BAB '): continue
            if lang == 'id' and bab_id.startswith('Chapter '): continue

            bab_title = bab_meta[bab_id]['title']
            has_content = len(items) > 0

            p_entry = doc.add_paragraph()
            p_entry.paragraph_format.first_line_indent = Pt(0)
            p_entry.paragraph_format.space_before = Pt(4)
            p_entry.paragraph_format.space_after  = Pt(4)

            # Tab stop kanan untuk nomor halaman
            pPr2 = p_entry._p.get_or_add_pPr()
            tabs2 = OxmlElement('w:tabs')
            t1 = OxmlElement('w:tab')
            t1.set(qn('w:val'), 'right')
            t1.set(qn('w:pos'), '9360')
            t1.set(qn('w:leader'), 'dot')   # ............. leader
            tabs2.append(t1)
            pPr2.append(tabs2)

            # Teks BAB
            r_bab = p_entry.add_run(f"{bab_id}   {bab_title}")
            r_bab.font.name  = 'Arial'
            r_bab.font.size  = Pt(11)
            r_bab.font.bold  = has_content
            r_bab.font.color.rgb = self.COLOR_BLACK if has_content else self.COLOR_GRAY

            # Tab + nomor halaman
            r_tab = p_entry.add_run("\t")
            r_pg  = p_entry.add_run(str(page_est))
            r_pg.font.name  = 'Arial'
            r_pg.font.size  = Pt(11)
            r_pg.font.bold  = has_content
            r_pg.font.color.rgb = self.COLOR_PRIMARY

            # Estimasi kasar: tiap BAB ~2 halaman jika ada konten
            page_est += (2 if has_content else 1)

        doc.add_page_break()

    # ─────────────────────────────────────────────
    # Heading Level Detector
    # ─────────────────────────────────────────────
    @staticmethod
    def _detect_heading_level(text: str, stored_level: int = 0) -> int:
        """
        Deteksi level heading. Robust terhadap OCR noise.

        Pola yang didukung (dengan/tanpa spasi ekstra dari OCR):
          "1. Pendahuluan"      -> H2  (level 2)
          "1.1 Sub Bab"        -> H3  (level 3)
          "1 .1 Sub Bab"       -> H3  (OCR: spasi di titik)
          "1.1.1 Detail"       -> H4  (level 4)
          "1.1. Detail"        -> H3  (trailing dot)
          "BAB I" / "BAB 1"    -> H2
        """
        # Prioritaskan stored_level dari DOCX style (paling akurat)
        if stored_level >= 2:
            return min(stored_level, 4)

        t = text.strip()

        # Pola: "BAB I", "BAB 1", "BAB II", dsb. → H2
        if re.match(r'^BAB\s+[IVXLC\d]+', t, re.IGNORECASE):
            return 2

        # Pola angka bernomor dengan titik — toleran terhadap OCR noise:
        normalized = re.sub(r'(\d)\s*\.\s*(?=\d)', r'\1.', t)

        m = re.match(r'^(\d+(?:\.\d+)*)\.?\s+\S', normalized)
        if m:
            number_part = m.group(1)
            dots = number_part.count('.')
            return min(2 + dots, 4)

        # Default: H2
        return 2

    # ─────────────────────────────────────────────
    # Main Builder
    # ─────────────────────────────────────────────
    def build_report(self, classified_data, original_filename, lang='id',
                     custom_product_name=None, custom_product_desc=None):
        doc = Document()

        self._set_fixed_margins(doc)
        self._set_fixed_styles(doc)

        # Use the explicit language from user selection
        doc_lang = lang
        bab_label_base = "Chapter" if doc_lang == 'en' else "BAB"
        self._add_header_footer(doc, bab_label=bab_label_base)

        # Ekstrak nama produk & deskripsi dari classified_data (skip generic titles)
        product_name = ""
        product_desc = ""
        generic_titles = {'user manual', 'manual book', 'buku manual', 'operating manual',
                          'instruction manual', 'table of contents', 'daftar isi', 'cover',
                          'introduction', 'pendahuluan', 'kata pengantar', 'preface'}
        # Cover-page specific terms that should NOT appear in BAB 1 content
        cover_only_titles = {'petunjuk pengguna', 'petunjuk pemakaian', 'user guide',
                             'owner manual', 'service manual', 'instruction manual',
                             'manual pengguna', 'panduan pengguna'}
        # Accept either BAB 1 or Chapter 1 for cover data
        cover_chapters = ('BAB 1', 'Chapter 1')
        cover_item_indices = set()  # Track indices of items used for cover page

        # Collect candidate headings (non-generic) from BAB 1
        heading_candidates = []
        for idx, item in enumerate(classified_data):
            if item.get('chapter_id') not in cover_chapters:
                continue
            item_type = item.get('type', '')
            text = item.get('normalized', '').strip()
            text_lower = text.lower() if text else ''

            # Pertimbangkan heading, title, DAN paragraph pendek sebagai kandidat nama produk
            if item_type in ('title', 'heading') or (item_type == 'paragraph' and len(text) < 50):
                # Mark generic cover titles for exclusion
                if text_lower in generic_titles or text_lower in cover_only_titles:
                    cover_item_indices.add(idx)
                    continue
                
                # Skip known brand names
                if 'elitech' in text_lower or 'technovision' in text_lower or text_lower.startswith('pt.'):
                    continue
                    
                if text and len(text) > 2:
                    heading_candidates.append((idx, text))
                    cover_item_indices.add(idx)

        # Pick product name: prefer heading with digits (model number), else first candidate
        product_name_idx = -1
        if heading_candidates:
            model_candidates = [(i, t) for i, t in heading_candidates if any(c.isdigit() for c in t)]
            if model_candidates:
                product_name_idx, product_name = model_candidates[0]
            else:
                product_name_idx, product_name = heading_candidates[0]

        # Description = next text element AFTER product name in document order
        # (short description like "Pengukur Panjang Badan Bayi dan Pengukur Tinggi Badan Dewasa")
        if product_name_idx >= 0:
            for idx, item in enumerate(classified_data):
                if idx <= product_name_idx:
                    continue
                if item.get('chapter_id') not in cover_chapters:
                    continue
                item_type = item.get('type', '')
                if item_type not in ('title', 'heading', 'paragraph'):
                    continue
                text = item.get('normalized', '').strip()
                text_lower = text.lower() if text else ''
                
                # Skip known brand names
                if 'elitech' in text_lower or 'technovision' in text_lower or text_lower.startswith('pt.'):
                    continue
                    
                if text and len(text) > 3 and text_lower not in generic_titles and text_lower not in cover_only_titles:
                    product_desc = text
                    cover_item_indices.add(idx)
                    break

        # Apply custom overrides if user has edited them
        if custom_product_name:
            product_name = custom_product_name
        if custom_product_desc:
            product_desc = custom_product_desc

        self._build_cover_page(doc, original_filename, product_name, product_desc, lang=doc_lang)

        # Kelompokkan per BAB (dilakukan lebih awal agar bisa dipakai TOC & konten)
        # Items that were used for the cover page are EXCLUDED from chapter content
        taxonomy_keys = list(BioBrain().taxonomy.keys())
        grouped = {k: [] for k in taxonomy_keys}
        for idx, item in enumerate(classified_data):
            # Skip items already used for cover page (by index or by flag)
            if idx in cover_item_indices or item.get('is_cover'):
                continue
            key = item['chapter_id']
            if key in grouped:
                grouped[key].append(item)
            else:
                # Fallback to first chapter based on lang
                fallback_key = "Chapter 1" if doc_lang == 'en' else "BAB 1"
                grouped[fallback_key].append(item)

        # Halaman Daftar Isi
        self._build_toc_page(doc, grouped, lang=doc_lang)

        # Bangun konten per BAB
        for bab_id, items in grouped.items():
            # Skip empty BAB 1-7 if documents use Chapter 1-7 (and vice versa)
            if doc_lang == 'en' and bab_id.startswith('BAB '): continue
            if doc_lang == 'id' and bab_id.startswith('Chapter '): continue

            bab_title = BioBrain().taxonomy[bab_id]["title"]
            self._add_chapter_header(doc, bab_id, bab_title)

            if not items:
                p = doc.add_paragraph("[ Tidak ada konten terdeteksi pada bab ini ]")
                p.italic = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                p.runs[0].font.color.rgb = self.COLOR_GRAY
                doc.add_page_break()
                continue

            current_heading_level = 2  # reset per chapter

            for item in items:
                content_type = item['type']
                text         = item['normalized']

                # ── Tabel dari DOCX tidak punya crop gambar, formatnya berwujud Markdown (Teks) ──
                # Jika dia 'table' tapi murni teks Markdown tanpa gambar crop (atau pakai fake PIL crop), alihkan ke pemroses teks
                resolved_crop = self._resolve_crop_path(item)
                if content_type == 'table' and (not resolved_crop or '_preview_only' in resolved_crop):
                    content_type = 'paragraph'

                if content_type in ('title', 'heading'):
                    # Tentukan level heading berdasarkan teks atau level tersimpan
                    stored_lvl = item.get('heading_level', 0)
                    lvl = self._detect_heading_level(text, stored_lvl)
                    current_heading_level = lvl  # track untuk indentasi body text
                    h = doc.add_heading(text, level=lvl)
                    h.paragraph_format.first_line_indent = Pt(0)

                elif content_type in ('figure', 'table'):
                    # ── Resolve crop image path with multiple fallback strategies ──
                    crop_local_val = self._resolve_crop_path(item)
                    logger.info(f"🖼️ {content_type}: crop_local={crop_local_val!r}, exists={os.path.exists(crop_local_val) if crop_local_val else 'N/A'}")
                    if crop_local_val and os.path.exists(crop_local_val):
                        # Check if the crop file is large enough to be an actual image
                        # Very small files (<5KB) are likely just cropped caption text
                        crop_file_size = os.path.getsize(crop_local_val)
                        logger.info(f"   📏 crop size: {crop_file_size} bytes")
                        if crop_file_size < 5000:  # Less than 5KB
                            # Too small — probably just caption text, render as paragraph
                            if text and text not in ("[TABLE DATA DETECTED]", "[FIGURE]", "[TABLE]"):
                                p = doc.add_paragraph(text)
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p.paragraph_format.first_line_indent = Pt(0)
                                p.runs[0].italic = True
                                p.runs[0].font.size = Pt(10)
                                p.runs[0].font.color.rgb = self.COLOR_GRAY
                            continue

                        # Gambar dengan border tipis via tabel 1-cell
                        try:
                            from docx.enum.text import WD_LINE_SPACING
                            tbl = doc.add_table(rows=1, cols=1)
                            tbl.style = 'Table Grid'
                            cell = tbl.cell(0, 0)
                            cell_para = cell.paragraphs[0]
                            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cell_para.paragraph_format.first_line_indent = Pt(0)
                            # FIX: remove exact line spacing that clips the image
                            cell_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            cell_para.paragraph_format.space_before = Pt(0)
                            cell_para.paragraph_format.space_after = Pt(0)
                            run_img = cell_para.add_run()
                            
                            # Cek lebar/resolusi untuk insert
                            try:
                                from PIL import Image
                                with Image.open(crop_local_val) as img_temp:
                                    w, h = img_temp.size
                                # Jika gambar terlalu kecil, gunakan fallback
                                if w < 50 and h < 50:
                                    doc.add_paragraph(f"[ Gambar terlalu kecil ]")
                                    continue
                            except Exception:
                                pass

                            run_img.add_picture(crop_local_val, width=Inches(4.5))
                            logger.info(f"   ✅ Picture added successfully: {os.path.basename(crop_local_val)}")
                        except Exception as e:
                            logger.error(f"Error adding picture: {e}")
                            doc.add_paragraph(f"[ Gagal load gambar: {e} ]")

                        # Caption
                        if text and text not in ("[TABLE DATA DETECTED]", "[FIGURE]", "[TABLE]"):
                            cap = doc.add_paragraph(text)
                            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cap.paragraph_format.space_before = Pt(4)
                            cap.paragraph_format.first_line_indent = Pt(0)
                            cap.runs[0].italic = True
                            cap.runs[0].font.size = Pt(9)
                            cap.runs[0].font.color.rgb = self.COLOR_GRAY

                        doc.add_paragraph()   # spasi setelah gambar

                    else:
                        logger.warning(f"   ❌ Could not resolve any crop path for {content_type}")
                        logger.warning(f"      item keys: {list(item.keys())}")
                        logger.warning(f"      crop_local: {item.get('crop_local')!r}")
                        logger.warning(f"      crop_url: {item.get('crop_url')!r}")
                        p = doc.add_paragraph(f"[ {content_type} — gambar tidak tersedia ]")
                        p.paragraph_format.first_line_indent = Pt(0)
                        p.runs[0].font.color.rgb = self.COLOR_GRAY

                else:  # Body text
                    lines = text.strip().split('\n')
                    is_md_table = False
                    if len(lines) >= 3 and '|' in lines[0] and '|' in lines[1] and ('-' in lines[1] or '=' in lines[1]):
                        is_md_table = True
                        
                    if is_md_table:
                        table_data = []
                        for line in lines:
                            line = line.strip()
                            if not line or '|-' in line or '-|' in line or '---|' in line:
                                continue
                            if line.startswith('|'): line = line[1:]
                            if line.endswith('|'): line = line[:-1]
                            cols = [c.strip() for c in line.split('|')]
                            table_data.append(cols)
                            
                        if table_data and len(table_data) > 0:
                            num_cols = max(len(row) for row in table_data)
                            tbl = doc.add_table(rows=len(table_data), cols=num_cols)
                            tbl.style = 'Table Grid'
                            for r_idx, row in enumerate(table_data):
                                for c_idx, val in enumerate(row):
                                    if c_idx < num_cols:
                                        c = tbl.cell(r_idx, c_idx)
                                        c.text = val
                                        # Bold header
                                        if r_idx == 0:
                                            for para in c.paragraphs:
                                                for run in para.runs:
                                                    run.font.bold = True
                            doc.add_paragraph() # Spacing
                            continue

                    p = doc.add_paragraph(text)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    # Indentasi body text sesuai level heading induknya
                    body_indent_map = {2: 0.0, 3: 0.25, 4: 0.5}
                    body_indent = body_indent_map.get(current_heading_level, 0.0)
                    if body_indent > 0:
                        p.paragraph_format.left_indent = Inches(body_indent)
                    p.paragraph_format.first_line_indent = Inches(0.35) if body_indent == 0 else Pt(0)

                    if item.get('has_typo'):
                        warn = doc.add_paragraph()
                        warn.paragraph_format.first_line_indent = Pt(0)
                        wr = warn.add_run("⚠ Possible OCR typos detected")
                        wr.font.size = Pt(8)
                        wr.font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)  # Orange


            doc.add_page_break()

        # Simpan — gunakan product_name untuk nama file jika tersedia
        if product_name:
            # Bersihkan dari karakter yang tidak valid untuk nama file
            safe_name = re.sub(r'[\\/*?:"<>|]', '', product_name).strip()
            safe_name = safe_name[:80]  # batas 80 karakter
        else:
            safe_name = Path(original_filename).stem

        word_filename = f"{safe_name}.docx"
        word_path     = os.path.join(self.base_path, word_filename)

        # Set document metadata
        try:
            doc.core_properties.title   = safe_name
            doc.core_properties.subject = "Manual Book"
            doc.core_properties.author  = "BioManual AI"
        except Exception:
            pass

        doc.save(word_path)
        logger.info(f"✓ Word file saved: {word_filename}")

        # ── Konversi ke PDF ──────────────────────────────────────
        pdf_filename = f"{safe_name}.pdf"
        pdf_path = os.path.join(self.base_path, pdf_filename)
        try:
            from docx2pdf import convert
            convert(word_path, pdf_path)
            logger.info(f"✓ PDF file saved: {pdf_filename}")
        except Exception as e:
            logger.warning(f"⚠️ PDF conversion failed: {e}")
            pdf_filename = None

        return {
            "word_file": word_filename,
            "pdf_file":  pdf_filename
        }
