"""
DIRECT READER — Text-based PDF & DOCX Extraction (No OCR!)
============================================================

Handles documents that already contain embedded text:
  - DOCX → python-docx (100% accurate, preserves styles)
  - PDF  → pdfplumber  (100% accurate for text-based PDFs)

Falls back gracefully when the document is scanned/image-only.

Updated: March 2026
"""

import os
import re
import cv2
import logging
import numpy as np
import json
import base64
from pathlib import Path
import pdfplumber
import PyPDF2
from docx import Document

logger = logging.getLogger("BioManual.DirectReader")

_WATERMARK_PATTERNS = [
    r'(?i)draft', r'(?i)confidential', r'(?i)strictly confidential',
    r'(?i)internal use', r'(?i)preview', r'(?i)watermark', 
    r'(?i)sample', r'(?i)copy', r'(?i)not for distribution'
]

def _is_watermark(text: str) -> bool:
    """Check if the text matches common watermark patterns."""
    if not text: return False
    clean = text.strip()
    if len(clean) < 3 or len(clean) > 30: return False
    for p in _WATERMARK_PATTERNS:
        if re.search(p, clean):
            return True
    return False


# ═══════════════════════════════════════════════════════════════
# 1. PDF TYPE DETECTION — is it text-based or scanned?
# ═══════════════════════════════════════════════════════════════

def is_text_pdf(pdf_path: str, sample_pages: int = 5) -> dict:
    """
    Detect whether a PDF contains real embedded text or is scanned (image-only).
    
    Returns:
        {
            "is_text_based": bool,
            "text_ratio": float,      # 0.0-1.0 — ratio of pages with text
            "avg_chars_per_page": int, # Average characters per page
            "total_pages": int,
            "sample_text": str,       # First ~500 chars for language detection
        }
    """
    try:
        import PyPDF2
        
        with open(pdf_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            total_pages = len(reader.pages)
            check_pages = min(total_pages, sample_pages)
            
            text_pages = 0
            total_chars = 0
            sample_text_parts = []
            
            for i in range(check_pages):
                try:
                    text = reader.pages[i].extract_text() or ""
                    text = text.strip()
                    chars = len(text)
                    total_chars += chars
                    
                    # A page is "text-based" if it has at least 50 chars
                    if chars > 50:
                        text_pages += 1
                    
                    if len(' '.join(sample_text_parts)) < 500:
                        sample_text_parts.append(text)
                except Exception:
                    pass
            
            text_ratio = text_pages / max(check_pages, 1)
            avg_chars = total_chars // max(check_pages, 1)
            
            result = {
                "is_text_based": text_ratio >= 0.5 and avg_chars >= 80,
                "text_ratio": round(text_ratio, 2),
                "avg_chars_per_page": avg_chars,
                "total_pages": total_pages,
                "sample_text": ' '.join(sample_text_parts)[:500],
            }
            
            logger.info(
                f"📄 PDF Type Detection: {'TEXT-BASED' if result['is_text_based'] else 'SCANNED'} "
                f"(ratio={result['text_ratio']}, avg_chars={result['avg_chars_per_page']}, "
                f"pages={total_pages})"
            )
            return result
    
    except Exception as e:
        logger.warning(f"PDF type detection failed: {e}")
        return {
            "is_text_based": False,
            "text_ratio": 0.0,
            "avg_chars_per_page": 0,
            "total_pages": 0,
            "sample_text": "",
        }


# ═══════════════════════════════════════════════════════════════
# 2. DOCX DIRECT READER — 100% accurate text extraction
# ═══════════════════════════════════════════════════════════════

def extract_docx_direct(docx_path: str, lang: str = 'id') -> list[dict]:
    """
    Extract content directly from DOCX using python-docx.
    """
    try:
        with open(docx_path, 'rb') as f:
            doc = Document(f)
    except Exception as e:
        logger.error(f"Failed to open DOCX: {e}")
        return []
    elements = []
    y_position = 0  # Simulated vertical position for ordering
    LINE_HEIGHT = 30  # Simulated line height in pixels
    
    logger.info(f"📝 DOCX Direct Reader: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    
    # ── Track which paragraphs are inside tables (to skip them) ──
    # python-docx iterates body elements in order, including tables
    table_set = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    table_set.add(id(para))
    
    # ── Iterate body elements in document order ──
    # doc.element.body contains all elements (paragraphs + tables) in order
    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        
        if tag == 'p':
            # This is a paragraph element — find matching paragraph object
            para = None
            for p in doc.paragraphs:
                if p._element is child:
                    para = p
                    break
            
            if para is None or id(para) in table_set:
                continue
            
            text = para.text.strip()
            
            # Filter out watermarks
            if _is_watermark(text):
                continue
                
            if not text:
                # Check for images in inline shapes
                inline_shapes = para._element.findall(
                    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
                )
                if inline_shapes:
                    elements.append({
                        "type": "figure",
                        "text": "[FIGURE]",
                        "confidence": 1.0,
                        "bbox": [0, y_position, 800, y_position + 200],
                    })
                    y_position += 200
                continue
            
            # Determine type: heading vs paragraph
            style_name = (para.style.name or "").lower()
            elem_type = "paragraph"
            
            # Heading detection strategies:
            # 1. Word heading style (Heading 1, Heading 2, etc.)
            if 'heading' in style_name or 'title' in style_name:
                elem_type = "heading"
            # 2. All caps short text (common manual heading pattern)
            elif text.isupper() and len(text) < 80:
                elem_type = "heading"
            # 3. Bold text that is short (common heading pattern)
            elif len(text) < 80:
                runs = para.runs
                if runs and all(run.bold for run in runs if run.text.strip()):
                    elem_type = "heading"
            # 4. Large font size (relative to body text)
            if elem_type == "paragraph" and len(text) < 100:
                runs = para.runs
                if runs:
                    font_sizes = [run.font.size for run in runs if run.font.size]
                    if font_sizes:
                        max_size = max(font_sizes)
                        # Pt(14) = 177800 EMUs — typical heading threshold
                        if max_size and max_size >= 177800:
                            elem_type = "heading"
            
            # Check if paragraph contains inline images/drawings along with text
            inline_shapes = para._element.findall(
                './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
            )
            
            # Estimate bbox based on text length
            text_height = max(LINE_HEIGHT, (len(text) // 80 + 1) * LINE_HEIGHT)
            bbox = [50, y_position, 750, y_position + text_height]
            
            elements.append({
                "type": elem_type,
                "text": text,
                "confidence": 1.0,
                "bbox": bbox,
            })
            y_position += text_height + 5
            
            # If there were also images in this paragraph, add them
            if inline_shapes:
                elements.append({
                    "type": "figure",
                    "text": "[FIGURE]",
                    "confidence": 1.0,
                    "bbox": [50, y_position, 750, y_position + 200],
                })
                y_position += 200
        
        elif tag == 'tbl':
            # This is a table element — find matching table object
            table = None
            for t in doc.tables:
                if t._element is child:
                    table = t
                    break
            
            if table is None:
                continue
            
            # Extract table data as 2D array
            table_data = []
            table_text_parts = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                    if cell_text:
                        table_text_parts.append(cell_text)
                table_data.append(row_data)
            
            if not table_text_parts:
                continue
            
            # Create markdown representation of table
            table_md = _table_to_markdown(table_data)
            
            # Change: User wants tables to be "cut" (cropped) not turned into text.
            # For Word, we can't easily crop yet, but we'll stop showing the text.
            table_height = max(100, len(table_data) * LINE_HEIGHT)
            elements.append({
                "type": "table",
                "text": "[TABLE]",
                "confidence": 1.0,
                "bbox": [50, y_position, 750, y_position + table_height],
                "table_data": table_data,
                "markdown_text": table_md  # Keep as metadata
            })
            y_position += table_height + 10
    
    # ── Extract embedded images ──
    # Check for images that are standalone (not already captured)
    _extract_docx_images(doc, docx_path, elements)
    
    # ── Merge consecutive elements of the same type ──
    merged_elements = []
    for el in elements:
        if not merged_elements:
            merged_elements.append(el)
            continue
        
        last_el = merged_elements[-1]
        
        # Merge if both are text elements of the same type (heading or paragraph)
        if last_el['type'] == el['type'] and el['type'] in ('heading', 'paragraph'):
            # Merge text with newline
            last_el['text'] = last_el['text'] + "\n" + el['text']
            
            # Combine bounding boxes (min x, min y, max x, max y)
            bb1 = last_el.get('bbox', [0, 0, 0, 0])
            bb2 = el.get('bbox', [0, 0, 0, 0])
            last_el['bbox'] = [
                min(bb1[0], bb2[0]),
                min(bb1[1], bb2[1]),
                max(bb1[2], bb2[2]),
                max(bb1[3], bb2[3])
            ]
        else:
            merged_elements.append(el)
            
    elements = merged_elements
    
    logger.info(
        f"✅ DOCX Direct Reader: {len(elements)} elements "
        f"({sum(1 for e in elements if e['type'] == 'heading')} headings, "
        f"{sum(1 for e in elements if e['type'] == 'paragraph')} paragraphs, "
        f"{sum(1 for e in elements if e['type'] == 'table')} tables, "
        f"{sum(1 for e in elements if e['type'] == 'figure')} figures)"
    )
    
    return elements


def _table_to_markdown(table_data: list[list[str]]) -> str:
    """Convert 2D table data to markdown table string."""
    if not table_data or not table_data[0]:
        return ""
    
    # Header row
    header = table_data[0]
    md = "| " + " | ".join(h or " " for h in header) + " |\n"
    md += "| " + " | ".join("---" for _ in header) + " |\n"
    
    # Data rows
    for row in table_data[1:]:
        # Pad row to match header length
        padded = row + [""] * (len(header) - len(row))
        md += "| " + " | ".join(c or " " for c in padded[:len(header)]) + " |\n"
    
    return md.strip()


def _extract_docx_images(doc, docx_path: str, elements: list):
    """
    Extract embedded images from DOCX and save them to output_results/.
    Updates figure elements with crop_url and crop_local paths.
    """
    try:
        import zipfile
        from PIL import Image
        
        backend_dir = os.path.dirname(os.path.abspath(__file__))
        output_dir = os.path.join(backend_dir, "output_results")
        os.makedirs(output_dir, exist_ok=True)
        
        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        
        # DOCX is a ZIP file — extract images from word/media/
        with zipfile.ZipFile(docx_path, 'r') as z:
            media_files = [f for f in z.namelist() if f.startswith('word/media/')]
            
            if not media_files:
                return
            
            img_idx = 0
            figure_elements = [e for e in elements if e['type'] == 'figure']
            
            for media_path in sorted(media_files):
                ext = os.path.splitext(media_path)[1].lower()
                if ext not in ('.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff', '.emf', '.wmf'):
                    continue
                
                try:
                    img_data = z.read(media_path)
                    
                    # Save to output_results
                    crop_fname = f"{base_name}_direct_img_{img_idx}.png"
                    crop_path = os.path.join(output_dir, crop_fname)
                    
                    # Convert to PNG if needed
                    if ext in ('.emf', '.wmf'):
                        # Skip vector formats — can't easily convert
                        continue
                    
                    with open(crop_path, 'wb') as f:
                        f.write(img_data)
                    
                    from urllib.parse import quote
                    crop_url = f"http://127.0.0.1:8000/output/{quote(crop_fname)}"
                    
                    # Match with figure elements
                    if img_idx < len(figure_elements):
                        figure_elements[img_idx]['crop_url'] = crop_url
                        figure_elements[img_idx]['crop_local'] = crop_path
                    
                    img_idx += 1
                
                except Exception as e:
                    logger.warning(f"Failed to extract image {media_path}: {e}")
        
        logger.info(f"📸 Extracted {img_idx} images from DOCX")
    
    except Exception as e:
        logger.warning(f"DOCX image extraction error: {e}")


# ═══════════════════════════════════════════════════════════════
# 3. PDF DIRECT READER — text extraction via pdfplumber
# ═══════════════════════════════════════════════════════════════

def extract_pdf_direct(pdf_path: str, lang: str = 'id') -> dict:
    """
    Extract content directly from a text-based PDF using pdfplumber.
    
    Returns dict per page:
    {
        "pages": [
            {
                "page_num": int,
                "elements": [...],     # Same format as vision_engine elements
                "page_image_path": str # Path to page image (for preview & figure crops)
            }
        ]
    }
    """
    import cv2
    import numpy as np
    
    backend_dir = os.path.dirname(os.path.abspath(__file__))
    output_dir = os.path.join(backend_dir, "output_results")
    os.makedirs(output_dir, exist_ok=True)
    
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    
    pages_result = []
    
    with pdfplumber.open(pdf_path) as pdf:
        total_pages = len(pdf.pages)
        logger.info(f"📄 PDF Direct Reader: {total_pages} pages")
        
        for page_idx, page in enumerate(pdf.pages):
            # ── Visual Column Detection (Render + CV) ──
            # This allows us to handle 2, 3, or 4 column layouts without interleaving text.
            col_boundaries = [0, page.width]
            try:
                # Render low-res for speed
                render = page.to_image(resolution=72)
                img_cv = cv2.cvtColor(np.array(render.original), cv2.COLOR_RGB2BGR)
                h_cv, w_cv = img_cv.shape[:2]
                
                # Grayscale + Threshold
                gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
                _, binary = cv2.threshold(gray, 220, 255, cv2.THRESH_BINARY)
                
                # Vertical projection (percentage of white pixels)
                roi = binary[int(h_cv*0.1):int(h_cv*0.9), :] # Ignore headers/footers
                white_ratio = np.mean(roi == 255, axis=0) # shape (w_cv,)
                
                # Smooth
                kernel = np.ones(max(3, w_cv // 60)) / max(3, w_cv // 60)
                white_smooth = np.convolve(white_ratio, kernel, mode='same')
                
                # Identify gaps (> 85% white is safer for icons)
                is_gap = white_smooth > 0.85
                min_gap_w = max(5, int(w_cv * 0.012))
                
                detected_gaps = []
                in_gap = False
                gs = 0
                for x in range(w_cv):
                    if is_gap[x] and not in_gap:
                        gs = x
                        in_gap = True
                    elif not is_gap[x] and in_gap:
                        gw = x - gs
                        if gw >= min_gap_w:
                            gc = gs + gw // 2
                            # Ignore edges (15% margin)
                            if w_cv * 0.15 < gc < w_cv * 0.85:
                                detected_gaps.append({'center': gc * (page.width / w_cv)})
                        in_gap = False
                
                # AI FALLBACK: If heuristic fails but page is wide, ask AI
                # This handles complex 3-column layouts with icons in gutters
                if len(detected_gaps) == 0 and page.width > 500:
                    try:
                        from openrouter_client import get_openrouter_client
                        client = get_openrouter_client()
                        if client and client.is_available:
                            img_ai = cv2.resize(img_cv, (1000, 1000)) if w_cv > 1000 else img_cv
                            _, buffer = cv2.imencode('.jpg', img_ai, [cv2.IMWRITE_JPEG_QUALITY, 85])
                            img_b64 = base64.b64encode(buffer).decode('utf-8')
                            vision_model = os.getenv("AI_VISION_MODEL", "google/gemini-2.0-flash-001")
                            prompt = "Analyze the vertical column layout. Ignore icons in gutters. How many columns? Return JSON: {\"columns\": 2, \"gap_centers\": [50.5]}"
                            response = client.call(prompt, image_base64=img_b64, timeout=12)
                            if response:
                                jm = re.search(r'\{.*\}', response, re.DOTALL)
                                if jm:
                                    ai_res = json.loads(jm.group())
                                    if ai_res.get("columns", 1) >= 2 and ai_res.get("gap_centers"):
                                        detected_gaps = [{'center': page.width * (gc/100.0)} for gc in ai_res["gap_centers"]]
                    except: pass

                if detected_gaps:
                    centers = sorted([g['center'] for g in detected_gaps])
                    col_boundaries = [0] + centers + [page.width]
                    logger.info(f"📊 Page {page_idx+1}: detected {len(col_boundaries)-1} columns")
            except Exception as e:
                logger.warning(f"Column detection failed for page {page_idx+1}: {e}")

            # ── Extract elements for EACH column ────────────
            page_elements = []
            for col_idx in range(len(col_boundaries) - 1):
                cx1, cx2 = col_boundaries[col_idx], col_boundaries[col_idx+1]
                
                # Overlap pad (15px) to prevent truncation of edge characters
                pad = 15
                crop_bbox = (max(0, cx1 - pad), 0, min(page.width, cx2 + pad), page.height)
                col_page = page.crop(crop_bbox)
                
                # 1. Words
                raw_words = col_page.extract_words(x_tolerance=3, y_tolerance=3, keep_blank_chars=False, extra_attrs=['fontname', 'size'])
                
                # Filter margins
                header_margin = min(60, page.height * 0.08)
                footer_margin = max(page.height - 60, page.height * 0.92)
                words = [w for w in raw_words if header_margin < ((w['top']+w['bottom'])/2) < footer_margin]
                
                if words:
                    lines = _group_words_into_lines(words, page.height)
                    paragraphs = _merge_lines_into_paragraphs(lines, page.width, page.height)
                    avg_fs = np.mean([p['avg_size'] for p in paragraphs if p.get('avg_size')]) if paragraphs else 11
                    for para in paragraphs:
                        etype = "paragraph"
                        fs = para.get('avg_size', avg_fs)
                        text = para['text'].strip()
                        if not text: continue
                        if fs > avg_fs * 1.2 and len(text) < 100: etype = "heading"
                        elif text.isupper() and len(text) < 80: etype = "heading"
                        elif para.get('is_bold') and len(text) < 80: etype = "heading"
                        
                        # Ignore watermarks
                        if _is_watermark(text):
                            continue
                            
                        page_elements.append({
                            "type": etype, "text": text, "confidence": 1.0, "bbox": para['bbox']
                        })

                # 2. Tables
                tables = col_page.extract_tables({"vertical_strategy": "lines", "horizontal_strategy": "lines"})
                table_finder = col_page.find_tables({"vertical_strategy": "lines", "horizontal_strategy": "lines"})
                if tables:
                    for t_idx, table_data in enumerate(tables):
                        if not table_data or not any(any(c for c in r if c) for r in table_data): continue
                        clean_table = [[(c or "").strip() for c in r] for r in table_data]
                        table_md = _table_to_markdown(clean_table)
                        tb = table_finder[t_idx].bbox if t_idx < len(table_finder) else [cx1, 0, cx2, 100]
                        bbox = [int(tb[0]), int(tb[1]), int(tb[2]), int(tb[3])]
                        # Remove overlapping paragraphs
                        page_elements = [e for e in page_elements if not _bbox_overlap(e['bbox'], bbox, threshold=0.5)]
                        
                        # Change: [TABLE] placeholder, actual crop handled below
                        page_elements.append({
                            "type": "table", "text": "[TABLE]", "confidence": 1.0, "bbox": bbox, 
                            "table_data": clean_table, "markdown_text": table_md
                        })

            if not page_elements:
                pages_result.append({"page_num": page_idx, "elements": [], "is_empty": True})
            else:
                # Sort elements top-to-bottom
                page_elements.sort(key=lambda x: (x['bbox'][1], x['bbox'][0]))
                pages_result.append({
                    "page_num": page_idx,
                    "elements": page_elements,
                    "is_empty": False
                })
            
            # ── Extract images ──
            if hasattr(page, 'images') and page.images:
                for img_info in page.images:
                    img_bbox = [
                        int(img_info.get('x0', 0)),
                        int(img_info.get('top', 0)),
                        int(img_info.get('x1', 100)),
                        int(img_info.get('bottom', 100)),
                    ]
                    # Only add if above a minimum size
                    img_w = img_bbox[2] - img_bbox[0]
                    img_h = img_bbox[3] - img_bbox[1]
                    if img_w > 50 and img_h > 50:
                        page_elements.append({
                            "type": "figure",
                            "text": "[FIGURE]",
                            "confidence": 1.0,
                            "bbox": img_bbox,
                        })
            
            # ── 3. Generate Crops for Visual Elements (Requirement: "dipotong") ──
            if any(e['type'] in ('table', 'figure') for e in page_elements):
                try:
                    # Render high-res for cropping
                    crop_render = page.to_image(resolution=200)
                    crop_img_cv = cv2.cvtColor(np.array(crop_render.original), cv2.COLOR_RGB2BGR)
                    ch, cw = crop_img_cv.shape[:2]
                    
                    # Scale factor between PDF units and pixel units
                    scale_x = cw / page.width
                    scale_y = ch / page.height
                    
                    for idx, e in enumerate(page_elements):
                        if e['type'] in ('table', 'figure'):
                            ex1, ey1, ex2, ey2 = e['bbox']
                            
                            # Pad slightly
                            pad = 10
                            px1, py1 = max(0, int((ex1 - pad) * scale_x)), max(0, int((ey1 - pad) * scale_y))
                            px2, py2 = min(cw, int((ex2 + pad) * scale_x)), min(ch, int((ey2 + pad) * scale_y))
                            
                            crop_visual = crop_img_cv[py1:py2, px1:px2]
                            if crop_visual.size > 0:
                                crop_fname = f"{base_name}_p{page_idx}_{e['type']}_{idx}.png"
                                crop_path = os.path.join(output_dir, crop_fname)
                                cv2.imwrite(crop_path, crop_visual)
                                
                                from urllib.parse import quote
                                e['crop_url'] = f"http://127.0.0.1:8000/output/{quote(crop_fname)}"
                                e['crop_local'] = crop_path
                                logger.info(f"📸 PDF Direct Crop: {crop_fname}")
                except Exception as ex:
                    logger.warning(f"Failed to generate crops for page {page_idx+1}: {ex}")

            # Sort elements by vertical position (reading order)
            page_elements.sort(key=lambda e: (e['bbox'][1], e['bbox'][0]))
            
            # Save a page preview too
            preview_fname = f"PREVIEW_{base_name}_p{page_idx}.jpg"
            preview_path = os.path.join(output_dir, preview_fname)
            try:
                page.to_image(resolution=100).save(preview_path)
            except: pass

            pages_result.append({
                "page_num": page_idx,
                "elements": page_elements,
                "is_empty": len(page_elements) == 0,
                "page_image_local": preview_path
            })
    
    logger.info(
        f"✅ PDF Direct Reader: {len(pages_result)} pages, "
        f"{sum(len(p['elements']) for p in pages_result)} total elements"
    )
    
    return {"pages": pages_result}


def _group_words_into_lines(words: list, page_height: float) -> list:
    """Group words into lines based on Y position proximity."""
    if not words:
        return []
    
    # Sort by Y then X
    sorted_words = sorted(words, key=lambda w: (round(w['top'], 1), w['x0']))
    
    lines = []
    current_line = {
        'words': [sorted_words[0]],
        'top': sorted_words[0]['top'],
        'bottom': sorted_words[0]['bottom'],
    }
    
    for word in sorted_words[1:]:
        # Same line if Y is close (within 3 units)
        if abs(word['top'] - current_line['top']) < 3:
            current_line['words'].append(word)
            current_line['bottom'] = max(current_line['bottom'], word['bottom'])
        else:
            lines.append(current_line)
            current_line = {
                'words': [word],
                'top': word['top'],
                'bottom': word['bottom'],
            }
    
    lines.append(current_line)
    return lines


def _merge_lines_into_paragraphs(lines: list, page_width: float, page_height: float) -> list:
    """Merge adjacent lines into paragraphs based on spacing and indentation."""
    if not lines:
        return []
    
    paragraphs = []
    
    current_para = {
        'lines': [lines[0]],
        'text_parts': [' '.join(w['text'] for w in lines[0]['words'])],
        'top': lines[0]['top'],
        'bottom': lines[0]['bottom'],
        'sizes': [w.get('size', 11) for w in lines[0]['words'] if w.get('size')],
        'fonts': [w.get('fontname', '') for w in lines[0]['words'] if w.get('fontname')],
    }
    
    for line in lines[1:]:
        line_text = ' '.join(w['text'] for w in line['words'])
        line_top = line['top']
        prev_bottom = current_para['bottom']
        
        # Gap between this line and previous paragraph bottom
        gap = line_top - prev_bottom
        
        # Average line height in current paragraph
        avg_line_h = (current_para['bottom'] - current_para['top']) / max(len(current_para['lines']), 1)
        if avg_line_h < 5:
            avg_line_h = 12  # Default
        
        # Merge if gap is small (< 1.5x line height = same paragraph)
        if gap < avg_line_h * 1.5 and gap >= -2:
            current_para['lines'].append(line)
            current_para['text_parts'].append(line_text)
            current_para['bottom'] = line['bottom']
            current_para['sizes'].extend(w.get('size', 11) for w in line['words'] if w.get('size'))
            current_para['fonts'].extend(w.get('fontname', '') for w in line['words'] if w.get('fontname'))
        else:
            # Finalize current paragraph
            _finalize_paragraph(current_para, paragraphs, page_width)
            
            current_para = {
                'lines': [line],
                'text_parts': [line_text],
                'top': line['top'],
                'bottom': line['bottom'],
                'sizes': [w.get('size', 11) for w in line['words'] if w.get('size')],
                'fonts': [w.get('fontname', '') for w in line['words'] if w.get('fontname')],
            }
    
    # Finalize last paragraph
    _finalize_paragraph(current_para, paragraphs, page_width)
    
    return paragraphs


def _finalize_paragraph(para_data: dict, paragraphs: list, page_width: float):
    """Finalize a paragraph data structure."""
    text = ' '.join(para_data['text_parts']).strip()
    if not text:
        return
    
    # Calculate bounding box
    all_words = []
    for line in para_data['lines']:
        all_words.extend(line['words'])
    
    if not all_words:
        return
    
    x0 = min(w['x0'] for w in all_words)
    x1 = max(w['x1'] for w in all_words)
    
    bbox = [int(x0), int(para_data['top']), int(x1), int(para_data['bottom'])]
    
    avg_size = np.mean(para_data['sizes']) if para_data['sizes'] else 11
    
    # Detect if bold font (heuristic: font name contains "Bold")
    is_bold = any('bold' in f.lower() for f in para_data['fonts'] if f)
    
    paragraphs.append({
        'text': text,
        'bbox': bbox,
        'avg_size': float(avg_size),
        'is_bold': is_bold,
    })


def _bbox_overlap(bbox1, bbox2, threshold=0.5) -> bool:
    """Check if two bboxes overlap significantly."""
    x1 = max(bbox1[0], bbox2[0])
    y1 = max(bbox1[1], bbox2[1])
    x2 = min(bbox1[2], bbox2[2])
    y2 = min(bbox1[3], bbox2[3])
    
    if x2 <= x1 or y2 <= y1:
        return False
    
    overlap_area = (x2 - x1) * (y2 - y1)
    area1 = max((bbox1[2] - bbox1[0]) * (bbox1[3] - bbox1[1]), 1)
    
    return overlap_area / area1 > threshold
